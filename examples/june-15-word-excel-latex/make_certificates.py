# Day 5 example 2 — 20 certificates in one go (openpyxl reads, python-docx writes)
# Run:  python make_certificates.py   ->  names.xlsx + certificates/ (20 files)
import os
from datetime import date
from openpyxl import Workbook, load_workbook
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

NAMES = ["Amy", "Ben", "Carla", "Diego", "Emma", "Felix", "Grace", "Hugo",
         "Iris", "Jack", "Kira", "Liam", "Mia", "Noah", "Olivia", "Pablo",
         "Quinn", "Rosa", "Sam", "Tina"]

# step 1 — make the name list (in real life the teacher already has this file)
wb = Workbook()
ws = wb.active
ws.append(["Name"])
for n in NAMES:
    ws.append([n])
wb.save("names.xlsx")

# step 2 — one certificate per name
os.makedirs("certificates", exist_ok=True)
names = [row[0].value for row in load_workbook("names.xlsx").active.iter_rows(min_row=2)]
for name in names:
    doc = Document()
    title = doc.add_heading("Certificate of Achievement", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"\nAwarded to\n\n{name}\n\nfor finishing Week 1 of the "
                    f"Vibe Coding class.\n\n{date.today():%B %d, %Y}")
    run.font.size = Pt(16)
    doc.save(f"certificates/{name.lower()}.docx")

print(f"done — {len(names)} certificates in certificates/")
