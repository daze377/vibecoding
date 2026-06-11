# Day 5 example 1 — a Word report written by code (python-docx)
# Run:  python make_report.py   ->  report.docx
from docx import Document

doc = Document()
doc.add_heading("Pet Survey Report", 0)
doc.add_paragraph("We asked 24 students which pet they like best. "
                  "Here is what we found.")

doc.add_heading("Results", level=1)
table = doc.add_table(rows=1, cols=2)
table.style = "Light Grid Accent 1"
table.rows[0].cells[0].text = "Pet"
table.rows[0].cells[1].text = "Votes"
for pet, votes in [("Dog", 10), ("Cat", 8), ("Hamster", 4), ("Fish", 2)]:
    row = table.add_row()
    row.cells[0].text = pet
    row.cells[1].text = str(votes)

doc.add_heading("What we learned", level=1)
doc.add_paragraph("Dogs win! 10 of 24 students picked a dog. "
                  "Cats came a close second with 8 votes.")

doc.save("report.docx")
print("report.docx created — open it in Word!")
